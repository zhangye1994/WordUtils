import time

from docx import Document
from docx.shared import Inches

from reader.Reader import Reader


class Writer:
    file_path = None
    doc = None

    def __init__(self, file_path):
        self.file_path = file_path
        self.doc = Document()

    def add_heading(self, text, level=1):
        """
        添加word段落

        :param text: 段落文本
        :return: 对象自身
        """
        self.doc.add_heading(text, level)
        return self

    def add_paragraph(self, text):
        """
        添加word段落

        :param text: 段落文本
        :return: 对象自身
        """
        self.doc.add_paragraph(text)
        return self

    def add_table(self, data=None, style='Table Grid'):
        """
        添加数据到表格，请注意，首行列数一定要最大，如果不是请填充空白数据
            *原因：不想在这种简单数据且不常见的场景处做过多的判断

        :param style: 表格样式
        :param data: 数据集
        :return: 对象自身
        """
        if data is None:
            data = [[]]
        rows = len(data)
        cols = 0 if len(data[0]) == 0 else len(data[0])
        table = self.doc.add_table(rows, cols, style)
        for row in range(rows):
            for col in range(cols):
                text = str(data[row][col]).replace("\r", "").replace("\x07", "")
                table.cell(row, col).text = text
        return self

    def set_width(self, tb_index=0, start=0, end=0, width=0):
        table = self.doc.tables[tb_index]
        columns = table.columns
        for index in range(start, end):
            columns[index].width = Inches(width)
        return self

    def set_height(self, tb_index=0, start=0, end=0, height=0):
        table = self.doc.tables[tb_index]
        rows = table.rows
        for index in range(start, end):
            rows[index].height = Inches(height)
        return self

    def save(self):
        """
        保存文档

        :return: None
        """
        if self.file_path is None or self.file_path is None:
            return
        self.doc.save(self.file_path)
